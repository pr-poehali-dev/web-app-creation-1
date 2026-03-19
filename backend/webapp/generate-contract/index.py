"""
Генерация форвардного контракта или договора на бартер по ГК РФ в формате DOCX.
v1.1
Принимает данные формы + данные пользователя, возвращает base64-файл для скачивания
и сохраняет в S3 для последующей загрузки в «Мои контракты».
"""
import json
import os
import io
import base64
import boto3
import psycopg2
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from templates import get_forward_contract_text, get_barter_contract_text


def get_db():
    return psycopg2.connect(os.environ["DATABASE_URL"])


def get_user(user_id: int) -> dict:
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        "SELECT id, email, first_name, last_name, middle_name, phone, user_type, "
        "company_name, inn, ogrnip, ogrn, legal_address, city, region, director_name "
        "FROM users WHERE id = %s",
        (user_id,)
    )
    row = cur.fetchone()
    conn.close()
    if not row:
        return {}
    keys = ["id", "email", "firstName", "lastName", "middleName", "phone", "userType",
            "companyName", "inn", "ogrnip", "ogrn", "legalAddress", "city", "region", "directorName"]
    return dict(zip(keys, row))


def build_docx(paragraphs: list) -> bytes:
    doc = Document()

    # Поля страницы
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(0.8)

    # Стили
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for item in paragraphs:
        kind = item[0]

        if kind == "spacer":
            doc.add_paragraph("")

        elif kind == "title":
            p = doc.add_paragraph(item[1])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(14)

        elif kind == "subtitle":
            p = doc.add_paragraph(item[1])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(12)

        elif kind == "city_date":
            p = doc.add_paragraph(item[1])
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif kind == "section":
            p = doc.add_paragraph(item[1])
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(12)

        elif kind == "text":
            p = doc.add_paragraph(item[1])
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(0)

        elif kind == "parties":
            _, seller_req, buyer_req = item
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            # убираем границы
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            def no_borders(tbl):
                tblPr = tbl._tbl.tblPr
                tblBorders = OxmlElement("w:tblBorders")
                for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    border = OxmlElement(f"w:{border_name}")
                    border.set(qn("w:val"), "none")
                    tblBorders.append(border)
                tblPr.append(tblBorders)

            no_borders(table)

            cells = table.rows[0].cells
            cells[0].text = "СТОРОНА 1 / ПОСТАВЩИК:\n\n" + seller_req
            cells[1].text = "СТОРОНА 2 / ПОКУПАТЕЛЬ:\n\n" + buyer_req

            for cell in cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def save_to_s3(docx_bytes: bytes, filename: str) -> str:
    s3 = boto3.client(
        "s3",
        endpoint_url="https://bucket.poehali.dev",
        aws_access_key_id=os.environ["AWS_ACCESS_KEY_ID"],
        aws_secret_access_key=os.environ["AWS_SECRET_ACCESS_KEY"],
    )
    key = f"contracts/{filename}"
    s3.put_object(
        Bucket="files",
        Key=key,
        Body=docx_bytes,
        ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    project_id = os.environ["AWS_ACCESS_KEY_ID"]
    return f"https://cdn.poehali.dev/projects/{project_id}/bucket/{key}"


def handler(event: dict, context) -> dict:
    """
    Генерация форвардного контракта или договора на бартер по ГК РФ.
    POST /generate-contract
    Body: { contractType, data (поля формы), buyerData (объект контрагента, опционально) }
    Возвращает: { docxBase64, docxUrl, filename }
    """
    cors = {
        "Access-Control-Allow-Origin": "*",
        "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
        "Access-Control-Allow-Headers": "Content-Type, X-User-Id",
    }

    if event.get("httpMethod") == "OPTIONS":
        return {"statusCode": 200, "headers": cors, "body": ""}

    user_id = event.get("headers", {}).get("x-user-id") or event.get("headers", {}).get("X-User-Id")
    if not user_id:
        return {"statusCode": 401, "headers": cors, "body": json.dumps({"error": "Unauthorized"})}

    body = json.loads(event.get("body") or "{}")
    data = body.get("data", {})
    contract_type = data.get("contractType", "forward")

    seller = get_user(int(user_id))
    if not seller:
        return {"statusCode": 404, "headers": cors, "body": json.dumps({"error": "User not found"})}

    # Данные покупателя/контрагента — или из body, или заглушка
    buyer = body.get("buyerData") or {
        "firstName": data.get("counterpartyName", ""),
        "lastName": "",
        "companyName": data.get("counterpartyCompany", ""),
        "inn": data.get("counterpartyInn", ""),
        "city": data.get("counterpartyCity", ""),
        "phone": data.get("counterpartyPhone", ""),
        "email": data.get("counterpartyEmail", ""),
        "userType": data.get("counterpartyType", "individual"),
    }

    today_str = date.today().strftime("%Y%m%d")
    if contract_type == "barter":
        contract_num = f"БД-{today_str}-{seller['id']}"
        paragraphs = get_barter_contract_text(
            {**data, "contractNumber": contract_num}, seller, buyer
        )
        filename = f"barter_{seller['id']}_{today_str}.docx"
    else:
        contract_num = f"ФК-{today_str}-{seller['id']}"
        paragraphs = get_forward_contract_text(
            {**data, "contractNumber": contract_num}, seller, buyer
        )
        filename = f"forward_{seller['id']}_{today_str}.docx"

    docx_bytes = build_docx(paragraphs)
    docx_b64 = base64.b64encode(docx_bytes).decode("utf-8")
    docx_url = save_to_s3(docx_bytes, filename)

    return {
        "statusCode": 200,
        "headers": {**cors, "Content-Type": "application/json"},
        "body": json.dumps({
            "success": True,
            "docxBase64": docx_b64,
            "docxUrl": docx_url,
            "filename": filename,
            "contractNumber": contract_num,
        }),
    }